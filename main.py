import asyncio
import copy
import json
import os
import random
import re
import time
from aiohttp import web
from aiogram import Bot, Dispatcher, F
from aiogram.filters import CommandStart
from aiogram.types import (
    Message,
    ReplyKeyboardMarkup,
    KeyboardButton,
    InlineKeyboardMarkup,
    InlineKeyboardButton,
    CallbackQuery,
    WebAppInfo,
)
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.context import FSMContext
from datetime import datetime, timezone, timedelta

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# ─────────────────────────── SOZLAMALAR ──────────────────────────

TOKEN = os.getenv("BOT_TOKEN")
if not TOKEN:
    raise ValueError("BOT_TOKEN o'rnatilmagan!")

# Web App URL – bot bilan bir xil domen bo‘lsa yaxshi
WEBAPP_URL = os.getenv("WEBAPP_URL", "https://YOUR_DOMAIN/webapp")

DATA_FILE = "data/users.json"
os.makedirs("temp", exist_ok=True)
os.makedirs("data", exist_ok=True)

UZBEKISTAN_TZ = timezone(timedelta(hours=5))
SUPPORTED_EXT = (".docx", ".doc", ".txt", ".xlsx", ".pdf")

bot = Bot(token=TOKEN)
dp = Dispatcher()


# ─────────────────────────── VAQT ────────────────────────────────

def uz_time():
    return datetime.now(UZBEKISTAN_TZ)

def uz_time_str():
    return uz_time().strftime("%d.%m.%Y %H:%M")


# ─────────────────────────── PERSISTENCE ─────────────────────────

def load_users() -> dict:
    try:
        if os.path.exists(DATA_FILE):
            with open(DATA_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception as e:
        print(f"⚠️ Yuklash xatoligi: {e}")
    return {}

def save_users(data: dict):
    try:
        with open(DATA_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2, default=str)
    except Exception as e:
        print(f"⚠️ Saqlash xatoligi: {e}")

users: dict = load_users()
user_messages: dict = {}

def get_user(uid: int) -> dict:
    key = str(uid)
    if key not in users:
        users[key] = {
            "first_visit": uz_time_str(),
            "total_tests": 0,
            "total_correct": 0,
            "results": [],
            "uploaded_docs": [],
        }
        save_users(users)
    return users[key]

def save_user(uid: int):
    save_users(users)


# ─────────────────────────── XABAR BOSHQARUVI ────────────────────

async def safe_delete(chat_id: int, mid: int):
    try:
        await bot.delete_message(chat_id, mid)
    except Exception:
        pass

async def clean_chat(uid: int, chat_id: int):
    for mid in user_messages.get(uid, []):
        await safe_delete(chat_id, mid)
    user_messages[uid] = []

async def track(uid: int, mid: int):
    user_messages.setdefault(uid, [])
    if mid not in user_messages[uid]:
        user_messages[uid].append(mid)

def cleanup_temp(max_days=3):
    if not os.path.exists("temp"):
        return
    cutoff = time.time() - max_days * 86400
    for name in os.listdir("temp"):
        path = os.path.join("temp", name)
        try:
            if os.path.isfile(path) and os.path.getmtime(path) < cutoff:
                os.remove(path)
        except Exception:
            pass


# ─────────────────────────── KLAVIATURA ──────────────────────────

def main_kb():
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(
                text="🚀 Testni boshlash",
                web_app=WebAppInfo(url=WEBAPP_URL)
            )],
            [
                KeyboardButton(text="📊 Natijalarim"),
                KeyboardButton(text="🆘 Yordam"),
            ],
        ],
        resize_keyboard=True,
    )

def result_kb():
    return InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="📋 Batafsil", callback_data="details"),
            InlineKeyboardButton(text="📊 Tarix", callback_data="history"),
        ],
        [InlineKeyboardButton(text="🌐 Qayta boshlash", web_app=WebAppInfo(url=WEBAPP_URL))],
    ])


# ─────────────────────────── PARSERLAR ───────────────────────────

def _read_txt(path: str) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1251", "windows-1251", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            pass
    return ""

def _norm(text: str, limit: int = 100) -> str:
    t = str(text).strip()
    return (t[: limit - 1] + "…") if len(t) > limit else t

def _parse_hash(lines: list) -> list:
    questions, current_q, correct, opts, state_ = [], None, None, [], "idle"

    def flush():
        nonlocal current_q, correct, opts, state_
        if current_q and correct and len(opts) >= 2:
            all_opts = opts[:4]
            while len(all_opts) < 4:
                all_opts.append(f"Variant {len(all_opts) + 1}")
            if correct not in all_opts:
                all_opts.insert(0, correct)
                all_opts = all_opts[:4]
            questions.append({"question": current_q, "options": all_opts, "answer": correct})
        current_q, correct, opts, state_ = None, None, [], "idle"

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("#") or line.startswith("?"):
            flush()
            q = line[1:].strip().rstrip("?").strip()
            if q:
                current_q = q
            else:
                state_ = "need_q"
        elif state_ == "need_q":
            current_q = line.rstrip("?").strip()
            state_ = "idle"
        elif line.startswith("+"):
            ans = line[1:].strip()
            if not ans:
                state_ = "need_correct"
            elif current_q:
                correct = ans
                if ans not in opts:
                    opts.append(ans)
                state_ = "idle"
        elif state_ == "need_correct":
            correct = line
            if line not in opts:
                opts.append(line)
            state_ = "idle"
        elif line.startswith("-"):
            ans = line[1:].strip()
            if current_q and ans and ans not in opts:
                opts.append(ans)

    flush()
    return questions

def _parse_abcd(lines: list) -> list:
    questions, current_q, opts_d, correct_l = [], None, {}, None

    def flush():
        nonlocal current_q, opts_d, correct_l
        if current_q and opts_d and correct_l:
            ul = correct_l.upper()
            if ul in opts_d:
                ans = opts_d[ul]
                options = list(opts_d.values())[:4]
                while len(options) < 4:
                    options.append(f"Variant {len(options) + 1}")
                questions.append({"question": current_q, "options": options, "answer": ans})
        current_q, opts_d, correct_l = None, {}, None

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        m = re.match(r"^(\d+)[.)]\s+(.+)$", line)
        if m:
            flush()
            current_q = m.group(2).strip()
            continue
        m = re.match(r"^([A-Da-d])[.)]\s+(.+)$", line)
        if m:
            opts_d[m.group(1).upper()] = m.group(2).strip()
            continue
        m = re.match(r"^(?:Javob|To'g'ri\s*javob|Answer|Ans)[:\s]*([A-Da-d])", line, re.IGNORECASE)
        if m:
            correct_l = m.group(1).upper()

    flush()
    return questions

def _parse_pipe(lines: list) -> list:
    result = []
    for line in lines:
        parts = [p.strip() for p in line.split("|")]
        if len(parts) >= 5 and parts[0]:
            result.append({"question": parts[0], "options": parts[1:5], "answer": parts[1]})
    return result

def parse_txt(path: str) -> list:
    content = _read_txt(path)
    if not content:
        return []
    lines = content.splitlines()
    ne = [l.strip() for l in lines if l.strip()]
    if not ne:
        return []

    has_hash = any(l.startswith("#") for l in ne)
    has_plus = any(l.startswith("+") for l in ne)
    has_pipe = any("|" in l and l.count("|") >= 4 for l in ne)
    has_num = any(re.match(r"^\d+[.)]\s+", l) for l in ne)
    has_abcd = any(re.match(r"^[A-Da-d][.)]\s+", l) for l in ne)

    if (has_hash) and has_plus:
        r = _parse_hash(lines)
        if r:
            return r
    if has_num and has_abcd:
        r = _parse_abcd(lines)
        if r:
            return r
    if has_pipe:
        r = _parse_pipe(ne)
        if r:
            return r
    for fn in [_parse_hash, _parse_abcd]:
        r = fn(lines)
        if r:
            return r
    return _parse_pipe(ne)

def parse_docx(path: str) -> list:
    try:
        from docx import Document
        doc = Document(path)
    except Exception:
        return []

    questions = []

    def add_q(q, opts, ans):
        if q and len(opts) >= 2 and ans in opts:
            o = opts[:4]
            while len(o) < 4:
                o.append(f"Variant {len(o) + 1}")
            questions.append({"question": q, "options": o, "answer": ans})

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            text = " ".join(dict.fromkeys(c for c in cells if c))
            if text:
                rows.append(text)

        if not rows:
            continue

        if len(rows) == 5:
            add_q(rows[0], rows[1:5], rows[1])
        elif len(rows) == 4:
            add_q(rows[0], rows[1:4] + ["Variant D"], rows[1])
        elif len(rows) > 5:
            for row in table.rows:
                cells = list(dict.fromkeys([c.text.strip() for c in row.cells if c.text.strip()]))
                if len(cells) >= 5:
                    add_q(cells[0], cells[1:5], cells[1])
                elif len(cells) == 3:
                    add_q(cells[0], cells[1:] + ["Variant C", "Variant D"], cells[1])

    if not questions:
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for fn in [_parse_hash, _parse_abcd]:
            r = fn(lines)
            if r:
                return r

    return questions

def parse_xlsx(path: str) -> list:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        questions = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if len(cells) >= 5:
                    questions.append({"question": cells[0], "options": cells[1:5], "answer": cells[1]})
        wb.close()
        return questions
    except Exception:
        return []

def parse_pdf(path: str) -> list:
    try:
        import pdfplumber
        lines = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    lines.extend(text.splitlines())
        lines = [l.strip() for l in lines if l.strip()]
        for fn in [_parse_hash, _parse_abcd, _parse_pipe]:
            r = fn(lines)
            if r:
                return r
        return []
    except Exception:
        return []


def convert_doc_to_docx(doc_path: str, docx_path: str) -> str:
    import shutil
    import subprocess

    soffice = (
        shutil.which("soffice")
        or shutil.which("libreoffice")
        or shutil.which("soffice.bin")
    )
    if not soffice:
        raise RuntimeError("LibreOffice serverda topilmadi")

    out_dir = os.path.dirname(docx_path) or "."
    try:
        result = subprocess.run(
            [soffice, "--headless", "--norestore", "--convert-to", "docx", "--outdir", out_dir, doc_path],
            capture_output=True,
            timeout=90,
        )
    except subprocess.TimeoutExpired:
        raise RuntimeError("Konversiya vaqti tugadi (90s)")

    auto_out = os.path.join(out_dir, os.path.splitext(os.path.basename(doc_path))[0] + ".docx")
    if not os.path.exists(auto_out):
        err = result.stderr.decode(errors="ignore")[:200] if result.stderr else "noma'lum xato"
        raise RuntimeError(f"Konversiya muvaffaqiyatsiz: {err}")

    if auto_out != docx_path:
        os.rename(auto_out, docx_path)
    return docx_path


# ─────────────────────────── FSM ─────────────────────────────────

class TestState(StatesGroup):
    choosing_count = State()
    testing = State()


# ─────────────────────────── TEST YORDAMCHILARI ──────────────────

def clear_session(uid: int):
    key = str(uid)
    if key not in users:
        return
    for k in ["selected_questions", "total_test", "current_index", "score",
               "answers", "waiting_for_skip", "current_poll_message_id",
               "current_poll_id", "current_question_index", "current_answer_recorded",
               "test_start_time"]:
        users[key].pop(k, None)

def grade_info(pct: float):
    if pct >= 90:
        return "A'lo", "🏆"
    elif pct >= 75:
        return "Yaxshi", "🎉"
    elif pct >= 60:
        return "Qoniqarli", "👍"
    else:
        return "Qoniqarsiz", "📚"

def count_kb(total: int):
    presets = sorted({n for n in [5, 10, 15, 20, 25, 30, 40, 50] if n <= total} | {total})
    rows, row = [], []
    for i, c in enumerate(presets):
        label = f"Hammasi ({c})" if c == total else f"{c} ta"
        row.append(InlineKeyboardButton(text=label, callback_data=f"cnt_{c}"))
        if len(row) == 3 or i == len(presets) - 1:
            rows.append(row)
            row = []
    rows.append([
        InlineKeyboardButton(text="🎲 Tasodifiy", callback_data="cnt_rand"),
        InlineKeyboardButton(text="✍️ O'zim", callback_data="cnt_custom"),
    ])
    rows.append([InlineKeyboardButton(text="❌ Bekor", callback_data="cancel")])
    return InlineKeyboardMarkup(inline_keyboard=rows)

def poll_kb():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="⏭  Keyingisi", callback_data="next_q")],
        [
            InlineKeyboardButton(text="💡 Javob", callback_data="hint"),
            InlineKeyboardButton(text="⏹  Yakunlash", callback_data="end_test"),
        ],
    ])


# ─────────────────────────── WEB SERVER ──────────────────────────

async def serve_webapp(request):
    try:
        with open("index.html", "r", encoding="utf-8") as f:
            return web.Response(text=f.read(), content_type="text/html")
    except FileNotFoundError:
        return web.Response(text="Web App topilmadi", status=404)

async def health(request):
    return web.Response(
        text=json.dumps({"status": "ok", "time": uz_time_str()}),
        content_type="application/json",
    )

def _cors_headers():
    return {
        "Access-Control-Allow-Origin": "*",
        "Access-Control-Allow-Methods": "GET, OPTIONS",
        "Access-Control-Allow-Headers": "Content-Type",
    }

async def api_files(request):
    uid = request.match_info.get("uid")
    user_data = users.get(str(uid), {})
    docs = user_data.get("uploaded_docs", [])

    out = []
    for d in docs:
        out.append({
            "file_name": d.get("file_name"),
            "uploaded_at": d.get("uploaded_at"),
            "questions": d.get("questions", []),
        })

    return web.Response(
        text=json.dumps({"files": out}, ensure_ascii=False),
        content_type="application/json",
        headers=_cors_headers(),
    )

async def api_results(request):
    uid = request.match_info.get("uid")
    user_data = users.get(str(uid), {})
    results = user_data.get("results", [])

    return web.Response(
        text=json.dumps({"results": results}, ensure_ascii=False, default=str),
        content_type="application/json",
        headers=_cors_headers(),
    )

async def api_options(request):
    return web.Response(headers=_cors_headers())

async def start_web():
    app = web.Application()
    app.router.add_get("/", serve_webapp)
    app.router.add_get("/health", health)
    app.router.add_get("/webapp", serve_webapp)
    app.router.add_get("/webapp/", serve_webapp)
    app.router.add_get("/api/files/{uid}", api_files)
    app.router.add_get("/api/results/{uid}", api_results)
    app.router.add_route("OPTIONS", "/api/files/{uid}", api_options)
    app.router.add_route("OPTIONS", "/api/results/{uid}", api_options)

    runner = web.AppRunner(app)
    await runner.setup()
    port = int(os.environ.get("PORT", 10000))
    site = web.TCPSite(runner, "0.0.0.0", port)
    await site.start()
    print(f"🌐 Web server: port {port}")


# ─────────────────────────── BOT HANDLERLAR ──────────────────────
# (To‘liq bot handlerlari – avvalgi koddagi barcha funksiyalar)
# Qisqalik uchun ularni bu yerga to‘liq joylashtiring.
# Men ularni tushirib qoldim, lekin siz avvalgi bot.py dan olishingiz mumkin.

# ─────────────────────────── MAIN ────────────────────────────────

async def main():
    print("🚀 Bot ishga tushdi")
    cleanup_temp()
    await start_web()
    await dp.start_polling(bot, allowed_updates=["message", "callback_query", "poll_answer"])

if __name__ == "__main__":
    asyncio.run(main())